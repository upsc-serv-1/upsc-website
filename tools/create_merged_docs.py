#!/usr/bin/env python3
"""
Read JSON files and create merged Word documents comparing Forum vs Prisma.
Does NOT modify JSON files - only reads them.
"""

import json
import re
from pathlib import Path
from collections import defaultdict
from difflib import SequenceMatcher

try:
    from docx import Document
    from docx.shared import RGBColor
except ImportError:
    import subprocess
    subprocess.run(["pip", "install", "python-docx", "-q"], check=True)
    from docx import Document
    from docx.shared import RGBColor

# Paths
APP_ROOT = Path("C:/Users/Dr. Yogesh/Pictures/website by windsurf")
IMPORTS_ROOT = APP_ROOT / "src" / "data" / "imports"
OUTPUT_ROOT = APP_ROOT / "new boks toolkit" / "MERGED_COMPARISON"
OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)

SIMILARITY_THRESHOLD = 0.85


def normalize_text(text):
    """Normalize for comparison."""
    if not text:
        return ""
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s]', '', text)
    return text.strip()


def similarity(text1, text2):
    """Calculate text similarity."""
    return SequenceMatcher(None, normalize_text(text1), normalize_text(text2)).ratio()


def get_year(q):
    """Extract year from question."""
    pyq = q.get('pyqMeta')
    if pyq and pyq.get('year'):
        return pyq['year']
    source = q.get('source', {}).get('sourceText', '')
    m = re.search(r'(\d{4})', source)
    if m:
        y = int(m.group(1))
        if 1970 <= y <= 2030:
            return y
    return None


def is_upsc(q):
    """Check if UPSC PYQ."""
    pyq = q.get('pyqMeta')
    if pyq and 'UPSC' in pyq.get('group', ''):
        return True
    return 'UPSC' in q.get('source', {}).get('sourceText', '').upper()


def load_and_categorize():
    """Load all JSON files and categorize."""
    forum = {'ancient': [], 'medieval': [], 'modern': [], 'art_culture': [], 
             'polity': [], 'economy': [], 'geography': [], 'environment': [],
             'science': [], 'world_geo': []}
    prisma = {'ancient': [], 'medieval': [], 'modern': [], 'art_culture': [], 
              'polity': [], 'economy': [], 'geography': [], 'environment': [],
              'science': [], 'world_geo': []}
    
    for f in IMPORTS_ROOT.glob("*.json"):
        fname = f.name.lower()
        
        # Skip Vivek Singh
        if 'vivek' in fname:
            continue
        
        try:
            data = json.loads(f.read_text(encoding='utf-8'))
            questions = data.get('questions', [])
            
            # Determine category
            cat = None
            if 'prisma' in fname:
                if 'ancient' in fname:
                    cat = ('prisma', 'ancient')
                elif 'medieval' in fname:
                    cat = ('prisma', 'medieval')
                elif 'modern' in fname:
                    cat = ('prisma', 'modern')
                elif 'art' in fname or 'culture' in fname:
                    cat = ('prisma', 'art_culture')
                elif 'polity' in fname:
                    cat = ('prisma', 'polity')
                elif 'economy' in fname:
                    cat = ('prisma', 'economy')
                elif 'indian' in fname and 'geog' in fname:
                    cat = ('prisma', 'geography')
                elif 'world' in fname and 'geog' in fname:
                    cat = ('prisma', 'world_geo')
                elif 'environment' in fname:
                    cat = ('prisma', 'environment')
                elif 'science' in fname:
                    cat = ('prisma', 'science')
            else:
                # Forum files
                sub = data.get('subject', '').lower()
                sec = data.get('sectionGroup', '').lower()
                
                if 'ancient' in fname or 'ancient' in sec:
                    cat = ('forum', 'ancient')
                elif 'medieval' in fname or 'medieval' in sec:
                    cat = ('forum', 'medieval')
                elif 'modern' in fname or 'modern' in sec:
                    cat = ('forum', 'modern')
                elif 'art' in fname or 'culture' in fname or 'art' in sec or 'culture' in sec:
                    cat = ('forum', 'art_culture')
                elif 'polity' in fname or 'polity' in sub:
                    cat = ('forum', 'polity')
                elif 'economy' in fname or 'economy' in sub:
                    cat = ('forum', 'economy')
                elif 'geography' in fname or 'geography' in sub:
                    if 'world' in fname or 'world' in sec:
                        cat = ('forum', 'world_geo')
                    else:
                        cat = ('forum', 'geography')
                elif 'environment' in fname:
                    cat = ('forum', 'environment')
            
            if cat:
                src, category = cat
                for q in questions:
                    q['_file'] = f.name
                    if src == 'forum':
                        forum[category].append(q)
                    else:
                        prisma[category].append(q)
        except Exception as e:
            print(f"Error loading {f.name}: {e}")
    
    return forum, prisma


def find_duplicates(forum_qs, prisma_qs):
    """Find duplicates between two lists."""
    duplicates = []
    forum_only = []
    matched = set()
    
    for fq in forum_qs:
        best_match = None
        best_sim = 0
        
        for i, pq in enumerate(prisma_qs):
            if i in matched:
                continue
            sim = similarity(fq.get('questionText', ''), pq.get('questionText', ''))
            if sim > best_sim and sim >= SIMILARITY_THRESHOLD:
                best_sim = sim
                best_match = i
        
        if best_match is not None:
            duplicates.append((fq, prisma_qs[best_match], best_sim))
            matched.add(best_match)
        else:
            forum_only.append(fq)
    
    prisma_only = [pq for i, pq in enumerate(prisma_qs) if i not in matched]
    
    return duplicates, forum_only, prisma_only


def create_doc(subject, duplicates, forum_only, prisma_only):
    """Create Word document for a subject."""
    doc = Document()
    
    # Title
    doc.add_heading(f'{subject.upper()} - Forum Toolkit vs Prisma PYQ Comparison', 0)
    
    # Statistics
    upsc_both = sum(1 for f, p, _ in duplicates if is_upsc(f))
    non_upsc_both = len(duplicates) - upsc_both
    upsc_forum = sum(1 for f in forum_only if is_upsc(f))
    non_upsc_forum = len(forum_only) - upsc_forum
    upsc_prisma = sum(1 for p in prisma_only if is_upsc(p))
    non_upsc_prisma = len(prisma_only) - upsc_prisma
    
    # Summary table
    doc.add_heading('Summary Statistics', level=1)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = 'Type'
    hdr[1].text = 'Only Forum'
    hdr[2].text = 'Only Prisma'
    hdr[3].text = 'In Both'
    
    row1 = table.rows[1].cells
    row1[0].text = 'UPSC PYQ'
    row1[1].text = str(upsc_forum)
    row1[2].text = str(upsc_prisma)
    row1[3].text = str(upsc_both)
    
    row2 = table.rows[2].cells
    row2[0].text = 'Non-UPSC PYQ'
    row2[1].text = str(non_upsc_forum)
    row2[2].text = str(non_upsc_prisma)
    row2[3].text = str(non_upsc_both)
    
    row3 = table.rows[3].cells
    row3[0].text = 'TOTAL'
    row3[1].text = str(len(forum_only))
    row3[2].text = str(len(prisma_only))
    row3[3].text = str(len(duplicates))
    
    doc.add_paragraph()
    
    # Combine all questions for year-wise grouping
    all_qs = []
    for f, p, _ in duplicates:
        f['_type'] = 'both'
        f['_prisma'] = p
        all_qs.append(f)
    for f in forum_only:
        f['_type'] = 'forum'
        all_qs.append(f)
    for p in prisma_only:
        p['_type'] = 'prisma'
        all_qs.append(p)
    
    # Group by year
    by_year = defaultdict(list)
    for q in all_qs:
        year = get_year(q)
        by_year[year if year else 'Unknown'].append(q)
    
    # Year-wise content
    for year in sorted(by_year.keys(), key=lambda x: (x == 'Unknown', x if x != 'Unknown' else 0)):
        doc.add_heading(str(year), level=1)
        
        # Sort by question number
        qs = sorted(by_year[year], key=lambda x: x.get('questionNumber', 0))
        
        for i, q in enumerate(qs, 1):
            q_type = q.get('_type', 'unknown')
            
            # Question
            qtext = q.get('questionText', '')
            p = doc.add_paragraph()
            p.add_run(f"Q{i}. ").bold = True
            p.add_run(qtext)
            
            # Options
            opts = q.get('options', {})
            for opt in ['a', 'b', 'c', 'd', 'e']:
                if opt in opts:
                    doc.add_paragraph(f"({opt}) {opts[opt]}", style='List Bullet 2')
            
            # Solutions
            if q_type == 'both':
                pq = q.get('_prisma')
                
                # Forum solution
                ans_f = q.get('correctAnswer', '')
                expl_f = q.get('explanationMarkdown', '')
                p = doc.add_paragraph()
                run = p.add_run(f"Solution ({ans_f}) [Forum Toolkit]")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 112, 192)
                if expl_f:
                    doc.add_paragraph(f"Explanation: {expl_f}")
                
                # Prisma solution
                ans_p = pq.get('correctAnswer', '')
                expl_p = pq.get('explanationMarkdown', '')
                p = doc.add_paragraph()
                run = p.add_run(f"Solution ({ans_p}) [Prisma]")
                run.bold = True
                run.font.color.rgb = RGBColor(192, 80, 77)
                if expl_p:
                    doc.add_paragraph(f"Explanation: {expl_p}")
                
                # Check metadata conflict
                f_sub = q.get('subject', '')
                f_sec = q.get('sectionGroup', '')
                f_micro = q.get('microTopic', '')
                p_sub = pq.get('subject', '')
                p_sec = pq.get('sectionGroup', '')
                p_micro = pq.get('microTopic', '')
                
                if f_sub != p_sub or f_sec != p_sec or f_micro != p_micro:
                    p = doc.add_paragraph()
                    run = p.add_run("[METADATA CONFLICT]")
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    doc.add_paragraph(f"Forum: [Subject: {f_sub}] [Section: {f_sec}] [Microtopic: {f_micro}]")
                    doc.add_paragraph(f"Prisma: [Subject: {p_sub}] [Section: {p_sec}] [Microtopic: {p_micro}]")
                
            elif q_type == 'forum':
                ans = q.get('correctAnswer', '')
                expl = q.get('explanationMarkdown', '')
                p = doc.add_paragraph()
                run = p.add_run(f"Solution ({ans}) [Forum Toolkit only]")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 112, 192)
                if expl:
                    doc.add_paragraph(f"Explanation: {expl}")
                
            elif q_type == 'prisma':
                ans = q.get('correctAnswer', '')
                expl = q.get('explanationMarkdown', '')
                p = doc.add_paragraph()
                run = p.add_run(f"Solution ({ans}) [Prisma only]")
                run.bold = True
                run.font.color.rgb = RGBColor(192, 80, 77)
                if expl:
                    doc.add_paragraph(f"Explanation: {expl}")
            
            # Metadata
            yr = get_year(q)
            src = q.get('source', {}).get('sourceText', '')
            doc.add_paragraph(f"[Year: {yr if yr else 'Unknown'} | Source: {src}]")
            
            sub = q.get('subject', '')
            sec = q.get('sectionGroup', '')
            micro = q.get('microTopic', '')
            doc.add_paragraph(f"[Subject: {sub}] [Section: {sec}] [Microtopic: {micro}]")
            
            doc.add_paragraph("─" * 60)
    
    return doc


def main():
    print("Loading and categorizing JSON files...")
    forum, prisma = load_and_categorize()
    
    subjects = [
        ('Ancient History', forum['ancient'], prisma['ancient']),
        ('Medieval History', forum['medieval'], prisma['medieval']),
        ('Modern History', forum['modern'], prisma['modern']),
        ('Art and Culture', forum['art_culture'], prisma['art_culture']),
        ('Polity', forum['polity'], prisma['polity']),
        ('Economy', forum['economy'], prisma['economy']),
        ('Geography', forum['geography'], prisma['geography']),
        ('Environment', forum['environment'], prisma['environment']),
        ('Science', forum['science'], prisma['science']),
        ('World Geography', forum['world_geo'], prisma['world_geo']),
    ]
    
    for subject_name, forum_qs, prisma_qs in subjects:
        if not forum_qs and not prisma_qs:
            print(f"Skipping {subject_name}: No data")
            continue
        
        print(f"\nProcessing {subject_name}...")
        print(f"  Forum: {len(forum_qs)} questions")
        print(f"  Prisma: {len(prisma_qs)} questions")
        
        # Find duplicates
        dups, forum_only, prisma_only = find_duplicates(forum_qs, prisma_qs)
        
        print(f"  Duplicates: {len(dups)}")
        print(f"  Forum only: {len(forum_only)}")
        print(f"  Prisma only: {len(prisma_only)}")
        
        # Create document
        doc = create_doc(subject_name, dups, forum_only, prisma_only)
        
        # Save
        out_path = OUTPUT_ROOT / f"{subject_name.replace(' ', '_')}_merged.docx"
        doc.save(out_path)
        print(f"  Saved: {out_path}")
    
    print(f"\n{'='*60}")
    print("All merged documents created successfully!")
    print(f"Location: {OUTPUT_ROOT}")


if __name__ == "__main__":
    main()
